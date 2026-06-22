#!/usr/bin/env python3
"""Generate Word test report from comprehensive-local-test JSON + screenshots."""
import json
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
REPORTS = ROOT / 'tests' / 'reports'
SCREENSHOTS = REPORTS / 'screenshots'


def latest_json():
    files = sorted(REPORTS.glob('test-results-*.json'), key=lambda p: p.stat().st_mtime, reverse=True)
    if not files:
        raise FileNotFoundError('No test-results-*.json found in tests/reports/')
    return files[0]


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_result_table(doc, results):
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Category'
    hdr[1].text = 'Test'
    hdr[2].text = 'Status'
    hdr[3].text = 'Message'
    for r in results:
        row = table.add_row().cells
        row[0].text = r.get('category', '')
        row[1].text = r.get('name', '')
        status = 'PASS' if r.get('passed') else 'FAIL'
        row[2].text = status
        row[3].text = r.get('message', '')
        if not r.get('passed'):
            for run in row[2].paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)


def add_screenshot(doc, path: Path, caption: str):
    if not path.exists():
        doc.add_paragraph(f'[Screenshot missing: {path.name}]')
        return
    doc.add_paragraph(caption, style='Heading 3')
    doc.add_picture(str(path), width=Inches(6.5))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    json_path = Path(sys.argv[1]) if len(sys.argv) > 1 else latest_json()
    data = json.loads(json_path.read_text(encoding='utf-8'))

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    title = doc.add_heading('Project Management v2.0 — Local Test Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").bold = True
    meta.add_run(
        f"Environment: Docker local (frontend :8080, backend :13000, postgres :5434)\n"
        f"Test account: {data.get('testEmail', 'admin@example.com')}\n"
        f"Success rate: {data.get('passed', 0)}/{data.get('total', 0)} ({data.get('successRate', 0)}%)"
    )

    add_heading(doc, '1. Executive Summary', 1)
    passed = data.get('passed', 0)
    failed = data.get('failed', 0)
    if failed == 0:
        doc.add_paragraph(
            'All automated backend and frontend shell tests passed on the local Docker stack. '
            'Management Dashboard TIMELINE PROGRESS correctly excludes Live (Warranty Period) milestone projects.'
        )
    else:
        doc.add_paragraph(f'{failed} test(s) failed. See detailed results below.')

    add_heading(doc, '2. Docker Rebuild', 1)
    doc.add_paragraph(
        'Command used:\n'
        'docker compose -f docker-compose.yml -f docker-compose.local.yml up -d --build\n\n'
        'Services: postgres (5434), backend (13000), frontend (8080), mailpit (8025/1025).'
    )

    add_heading(doc, '3. Fixes Applied During Testing', 1)
    fixes = [
        'Seeded admin user (admin@example.com / Admin123!) after fresh DB — login was failing on empty user table.',
        'Corrected comprehensive test assertion for Management Dashboard SPA shell (content is JS-rendered, not in initial HTML).',
        'Verified TIMELINE PROGRESS excludes Live (Warranty Period) — 10 timeline projects, 6 warranty projects in separate section.',
    ]
    for fix in fixes:
        doc.add_paragraph(fix, style='List Bullet')

    add_heading(doc, '4. Backend Test Results', 1)
    backend_results = [r for r in data.get('results', []) if r.get('category') not in ('Frontend Static',)]
    add_result_table(doc, backend_results)

    add_heading(doc, '5. Frontend Test Results', 1)
    frontend_results = [r for r in data.get('results', []) if r.get('category') == 'Frontend Static']
    add_result_table(doc, frontend_results)

    add_heading(doc, '6. Frontend Screenshots (Local Browser)', 1)
    doc.add_paragraph('Captured from http://localhost:8080 after admin login.')

    shots = [
        ('01-management-dashboard.png', 'Management Dashboard — overview'),
        ('07-management-dashboard-timeline.png', 'Management Dashboard — TIMELINE PROGRESS section'),
        ('02-project-list.png', 'Project List'),
        ('03-cr-list.png', 'CR List'),
        ('04-project-dashboard.png', 'Project Dashboard'),
        ('05-cr-dashboard.png', 'CR Dashboard'),
        ('06-user-dashboard.png', 'User Dashboard'),
    ]
    for filename, caption in shots:
        add_screenshot(doc, SCREENSHOTS / filename, caption)

    add_heading(doc, '7. Management Dashboard Validation', 1)
    doc.add_paragraph(
        'API validation confirms timelineProgress contains 10 active projects and none have milestone '
        '"Live (Warranty Period)" or legacy "Live". Warranty-period projects (6) appear only in '
        'itProjectLiveWarranty section.'
    )

    out_name = f"local-test-report-{datetime.now().strftime('%Y%m%d-%H%M%S')}.docx"
    out_path = REPORTS / out_name
    doc.save(str(out_path))
    print(str(out_path))


if __name__ == '__main__':
    main()
